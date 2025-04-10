# App.py creates the flask app, and renders html templates at #
# directories and handles function executions #

# Import libraries #
import flask
import datetime
import docx
import os
import time

# Import python files
from database import *

# Create app instance #
app = flask.Flask(__name__)
# Secret Key and configuration #
app.secret_key = "Secret_Key"
app.config.update(
    UPLOAD_FOLDER="images",
    TIME_SHEETS="timesheets",
    PAY_SHEETS="paysheets"
)

# Function which checks if the user has logged in based on their session data #
def checkLogged():
    try:
        userID = int(flask.session.get("UserID"))
        userType = flask.session.get("UserType")
        return True
    except:
        return False


# At the web application index render the index page #
@app.route("/", methods=["GET", "POST"])
def index():
    if checkLogged():
        return flask.redirect(flask.url_for("portal"))

    if flask.request.method == "GET":
        return flask.render_template("index.html")

    email = flask.request.form.get("email")
    password = flask.request.form.get("password")

    if email and password:
        SQL = """SELECT * FROM Accounts WHERE account_email = ?"""
        account = EXECUTE(SQL, values=(email,), getResult=True)
        if account:
            if account[0][1] == "owner":
                SQL = """SELECT * FROM BusinessOwner WHERE owner_email = ?"""
                owner = EXECUTE(SQL, values=(email,), getResult=True)
                if owner:
                    if password == owner[0][3]:
                        flask.session["UserType"] = "owner"
                        flask.session["UserID"] = owner[0][0]
                        return flask.redirect(flask.url_for("portal"))
                    else:
                        flask.flash("Incorrect Password!")
                        return flask.render_template("index.html")
                else:
                    return flask.render_template("index.html")

            elif account[0][1] == "manager":
                SQL = """SELECT * FROM Managers WHERE manager_email = ?"""
                manager = EXECUTE(SQL, values=(email,), getResult=True)
                if manager:
                    if password == manager[0][3]:
                        flask.session["UserType"] = "manager"
                        flask.session["UserID"] = manager[0][0]
                        return flask.redirect(flask.url_for("portal"))
                    else:
                        flask.flash("Incorrect Password!")
                        return flask.render_template("index.html")
                else:
                    return flask.render_template("index.html")

            elif account[0][1] == "employee":
                SQL = """SELECT * FROM Employees WHERE employee_email = ?"""
                employee = EXECUTE(SQL, values=(email,), getResult=True)
                if employee:
                    if password == employee[0][3]:
                        flask.session["UserType"] = "employee"
                        flask.session["UserID"] = employee[0][0]
                        return flask.redirect(flask.url_for("portal"))
                    else:
                        flask.flash("Incorrect Password!")
                        return flask.render_template("index.html")
                else:
                    return flask.render_template("index.html")

            else:
                return flask.render_template("index.html")
        else:
            flask.flash("Account Does Not Exist!")
            return flask.render_template("index.html")
    else:
        flask.flash("Please Fill All Sections!")
        return flask.render_template("index.html")

# Portal page for accessing each page #
@app.route("/portal")
def portal():
    if checkLogged() == False:
        return flask.redirect(flask.url_for("index"))

    userID = flask.session.get("UserID")
    userType = flask.session.get("UserType")

    SQL = """SELECT * FROM Business"""
    businessName = EXECUTE(SQL, getResult=True)[0][0]
    if userType == "owner":
        SQL = """SELECT * FROM BusinessOwner WHERE owner_id = ?"""
        userName = EXECUTE(SQL, values=(userID,), getResult=True)[0][1]
    elif userType == "manager":
        SQL = """SELECT * FROM Managers WHERE manager_id = ?"""
        userName = EXECUTE(SQL, values=(userID,), getResult=True)[0][1]
    elif userType == "employee":
        SQL = """SELECT * FROM Employees WHERE employee_id = ?"""
        userName = EXECUTE(SQL, values=(userID,), getResult=True)[0][1]
    return flask.render_template("portal.html", businessname=businessName, usertype=userType.capitalize(), username=userName)

# Shift scheduling route, implements the designed shift scheduling feature #
@app.route("/portal/shiftscheduling", methods=["GET", "POST"])
def shiftscheduling():
    if checkLogged() == False:
        return flask.redirect(flask.url_for("index"))

    SQL = """SELECT * FROM Business"""
    businessName = EXECUTE(SQL, getResult=True)[0][0]
    userID = flask.session.get("UserID")
    userType = flask.session.get("UserType")

    if userType == "owner":
        return flask.redirect(flask.url_for("portal"))

    elif userType == "manager":
        employees = []
        SQL = """SELECT * FROM Employees"""
        foundEmployees = EXECUTE(SQL, getResult=True)
        employeeDictionary = {}
        for employee in foundEmployees:
            employees.append(employee)
            employeeDictionary[employee[0]] = employee[1]


        shifts = [[],[],[],[],[],[],[]]
        SQL = """SELECT * FROM ShiftTemplate"""
        foundShifts = EXECUTE(SQL, getResult=True)

        for shift in foundShifts:
            shifts[shift[1]].append(shift)

        weekDates = []
        today = datetime.datetime.now()
        currentPosition = today.weekday()
        startDate = today - datetime.timedelta(days=currentPosition)
        endDate = today + datetime.timedelta(days=(7-currentPosition))
        while startDate != endDate:
            weekDates.append(startDate.strftime('%Y-%m-%d'))
            startDate = startDate + datetime.timedelta(days=1)

        actualShifts = []

        SQL = """SELECT * FROM Shifts WHERE date = ?"""
        for date in weekDates:
            foundActualShifts = EXECUTE(SQL, values=(date,), getResult=True)
            actualShift = []
            for foundActualShift in foundActualShifts:
                updatedEmployee = employeeDictionary.get(foundActualShift[1])
                updatedFoundActualShift = (foundActualShift[0], updatedEmployee, foundActualShift[3], foundActualShift[4], foundActualShift[5])
                actualShift.append(updatedFoundActualShift)
            actualShifts.append(actualShift)



        if flask.request.method == "GET":
            return flask.render_template("shiftscheduling.html", usertype=userType, shifts=shifts, employees=employees, businessname=businessName, actualshifts=actualShifts, leftbar=True, rightbar=True)

        action = flask.request.form.get("action")

        if action == "add":
            day = flask.request.form.get("day")
            starttime = flask.request.form.get("starttime")
            endtime = flask.request.form.get("endtime")
            employee = flask.request.form.get("employee")

            if day and starttime and endtime and employee:
                try:
                    starttime = datetime.datetime.strptime(starttime, "%H:%M").time()
                    endtime = datetime.datetime.strptime(endtime, "%H:%M").time()

                    if endtime <= starttime:
                        flask.flash("End Time Should Be After Start Time!")
                        return flask.render_template("shiftscheduling.html", usertype=userType, shifts=shifts, employees=employees, businessname=businessName, actualshifts=actualShifts, leftbar=True, rightbar=True)

                    starttime = starttime.strftime("%H:%M")
                    endtime = endtime.strftime("%H:%M")
                except:
                    flask.flash("Error: Please enter valid times (HH:MM in 24 hour format)!")
                    return flask.render_template("shiftscheduling.html", usertype=userType, shifts=shifts, employees=employees, businessname=businessName, actualshifts=actualShifts, leftbar=True, rightbar=True)

                SQL = """INSERT INTO ShiftTemplate (day, starttime, endtime, preffered_employee)
                            VALUES (?, ?, ?, ?)"""
                EXECUTE(SQL, values=(day, starttime, endtime, employee))
                return flask.redirect(flask.url_for("shiftscheduling"))
            else:
                flask.flash("Please Fill All Action Fields!")
                return flask.render_template("shiftscheduling.html", usertype=userType, shifts=shifts, employees=employees, businessname=businessName, actualshifts=actualShifts, leftbar=True, rightbar=True)

        elif action == "delete":
            shift = int(flask.request.form.get("shift"))
            if shift == 0:
                return flask.render_template("shiftscheduling.html", usertype=userType, shifts=shifts, employees=employees, businessname=businessName, actualshifts=actualShifts, leftbar=True, rightbar=True)

            SQL = """DELETE FROM ShiftTemplate
                        WHERE id = ?"""
            try:
                EXECUTE(SQL, values=(shift,))
                return flask.redirect(flask.url_for("shiftscheduling"))
            except:
                flask.flash("Shift Does Not Exist!")
                return flask.render_template("shiftscheduling.html", usertype=userType, shifts=shifts, employees=employees, businessname=businessName, actualshifts=actualShifts, leftbar=True, rightbar=True)

        elif action == "auto":
            shiftsGiven = []
            for employee in employees:
                shiftsGiven.append(0)
            print(shiftsGiven)

            currentDay = datetime.datetime.now().date()
            endDay = currentDay + datetime.timedelta(days=7)
            while currentDay < endDay:
                for shift in foundShifts:
                    if shift[1] == currentDay.weekday():
                        SQL1 = """SELECT * FROM Shifts
                                    WHERE employee_id = ?
                                    AND date = ?
                                    AND starttime = ?
                                    AND endtime = ?"""
                        SQL2 = """INSERT INTO Shifts (employee_id, date, starttime, endtime, accepted)
                                    VALUES (?, ?, ?, ?, ?)"""
                        if shift[4] != 0:
                            print("preffered")
                            isAssigned = EXECUTE(SQL1, values=(shift[4], currentDay.strftime('%Y-%m-%d'), shift[2], shift[3]), getResult=True)
                            if isAssigned:
                                continue
                            else:
                                EXECUTE(SQL2, values=(shift[4], currentDay.strftime('%Y-%m-%d'), shift[2], shift[3], 0))
                                shiftsGiven[shift[4]-1] += 1
                        else:
                            print("no preffered")
                            minVal = min(shiftsGiven)
                            minIndex = shiftsGiven.index(minVal)
                            isAssigned = EXECUTE(SQL1, values=(minIndex+1, currentDay.strftime('%Y-%m-%d'), shift[2], shift[3]), getResult=True)
                            if isAssigned:
                                continue
                            else:
                                EXECUTE(SQL2, values=(minIndex+1, currentDay.strftime('%Y-%m-%d'), shift[2], shift[3], 0))
                                shiftsGiven[minIndex] += 1
                    print(shiftsGiven)
                currentDay = currentDay + datetime.timedelta(days=1)
            return flask.redirect(flask.url_for("shiftscheduling"))

        elif action == "manualAdd":
            date = flask.request.form.get("date")
            starttime = flask.request.form.get("manualStarttime")
            endtime = flask.request.form.get("manualEndtime")
            employee = flask.request.form.get("manualEmployee")
            if date and starttime and endtime and employee:
                starttime = datetime.datetime.strptime(starttime, "%H:%M").time()
                endtime = datetime.datetime.strptime(endtime, "%H:%M").time()

                if endtime <= starttime:
                    flask.flash("End Time Should Be After Start Time!")
                    return flask.render_template("shiftscheduling.html", usertype=userType, shifts=shifts, employees=employees, businessname=businessName, actualshifts=actualShifts, leftbar=True, rightbar=True)

                starttime = starttime.strftime("%H:%M")
                endtime = endtime.strftime("%H:%M")

                SQL = """INSERT INTO SHIFTS (employee_id, date, starttime, endtime, accepted)
                            VALUES (?, ?, ?, ?, ?)"""
                EXECUTE(SQL, values=(employee, date, starttime, endtime, 0))
                return flask.redirect(flask.url_for("shiftscheduling"))
            else:
                flask.flash("Please Fill All Fields!")
                return flask.render_template("shiftscheduling.html", usertype=userType, shifts=shifts, employees=employees, businessname=businessName, actualshifts=actualShifts, leftbar=True, rightbar=True)

        elif action == "actualDelete":
            shiftID = flask.request.form.get("shiftID")
            SQL = """DELETE FROM Shifts
                        WHERE shift_id = ?"""
            try:
                EXECUTE(SQL, values=(shiftID,))
                return flask.redirect(flask.url_for("shiftscheduling"))
            except:
                flask.flash("Shift Does Not Exist!")
                return flask.render_template("shiftscheduling.html", usertype=userType, shifts=shifts, employees=employees, businessname=businessName, actualshifts=actualShifts, leftbar=True, rightbar=True)
        else:
            return flask.render_template("shiftscheduling.html", usertype=userType, shifts=shifts, employees=employees, businessname=businessName, actualshifts=actualShifts, leftbar=True, rightbar=True)

    elif userType == "employee":

        SQL = """SELECT * FROM Shifts
                    WHERE employee_id = ?
                    AND accepted = 0"""
        requestedShifts = EXECUTE(SQL, values=(userID,), getResult=True)

        currentDay = datetime.datetime.now()
        currentPosition = currentDay.weekday()
        startDate = currentDay - datetime.timedelta(days=currentPosition)
        endDate = currentDay + datetime.timedelta(days=(6-currentPosition))
        startDate = startDate.date()
        endDate = endDate.date()
        employeeShifts = []
        while startDate <= endDate:
            SQL = """SELECT * FROM Shifts
                        WHERE employee_id = ?
                        AND date = ?
                        AND accepted = 1"""
            foundShifts = EXECUTE(SQL, values=(userID, startDate), getResult=True)
            employeeShifts.append(foundShifts)
            startDate = startDate + datetime.timedelta(days=1)

        if flask.request.method == "GET":
            return flask.render_template("shiftscheduling.html", usertype=userType, businessname=businessName, requestedshifts=requestedShifts, employeeShifts=employeeShifts, rightbar=True)

        action = flask.request.form.get("requestedShiftAction")
        if action == "accept":
            shiftID = flask.request.form.get("shiftID")
            SQL = """UPDATE Shifts
                        SET accepted = 1
                        WHERE shift_id = ?"""
            EXECUTE(SQL, values=(shiftID,))
            return flask.redirect(flask.url_for("shiftscheduling"))
        elif action == "decline":
            shiftID = flask.request.form.get("shiftID")
            SQL = """SELECT * FROM Employees"""
            foundEmployees = EXECUTE(SQL, getResult=True)
            finalEmployee = foundEmployees[-1]
            if userID == finalEmployee[0]:
                nextEmployee = 1
            else:
                nextEmployee = userID + 1
            SQL = """UPDATE Shifts
                        SET employee_id = ?
                        WHERE shift_id = ?"""
            EXECUTE(SQL, values=(nextEmployee, shiftID))
            return flask.redirect(flask.url_for("shiftscheduling"))
        else:
            return flask.render_template("shiftscheduling.html", usertype=userType, businessname=businessName, requestedshifts=requestedShifts, rightbar=True)

    else:
        return flask.redirect(flask.url_for("portal"))

# Chat route, implements the designed communication feature from the research #
@app.route("/portal/chat", methods=["GET", "POST"])
def chat():
    if checkLogged() == False:
        return flask.redirect(flask.url_for("index"))

    SQL = """SELECT * FROM Business"""
    businessName = EXECUTE(SQL, getResult=True)[0][0]
    userID = flask.session.get("UserID")
    userType = flask.session.get("UserType")
    SQL = """SELECT * FROM Accounts
                WHERE account_type = ?
                AND type_id = ?"""
    accountID = (EXECUTE(SQL, values=(userType, userID), getResult=True))[0][0]

    SQL = """SELECT * FROM Accounts"""
    foundAccounts = EXECUTE(SQL, getResult=True)
    accountDictionary = {}
    for account in foundAccounts:
        if account[1] == "owner":
            SQL = """SELECT * FROM BusinessOwner
                        WHERE owner_id = ?"""
            foundName = EXECUTE(SQL, values=(account[3],), getResult=True)
            accountDictionary[account[0]] = foundName[0][1]
        elif account[1] == "manager":
            SQL = """SELECT * FROM Managers
                        WHERE manager_id = ?"""
            foundName = EXECUTE(SQL, values=(account[3],), getResult=True)
            accountDictionary[account[0]] = foundName[0][1]
        else:
            SQL = """SELECT * FROM Employees
                        WHERE employee_id = ?"""
            foundName = EXECUTE(SQL, values=(account[3],), getResult=True)
            accountDictionary[account[0]] = foundName[0][1]

    SQL = """SELECT * FROM GroupMessages"""
    foundMessages = EXECUTE(SQL, getResult=True)
    foundReplies = []
    updatedMessages = []
    updatedReplies = []
    for message in foundMessages:
        if message[2] == None:
            updatedMessage = (message[0], accountDictionary.get(message[1]), message[1], message[2], message[3])
            updatedMessages.append(updatedMessage)
        SQL = """SELECT * FROM GroupMessages
                    WHERE reply_id = ?"""
        currentReplies = EXECUTE(SQL, values=(message[0],), getResult=True)
        currentUpdatedReplies = []
        for reply in currentReplies:
            updatedReply = (reply[0], accountDictionary.get(reply[1]), reply[1], reply[2], reply[3])
            currentUpdatedReplies.append(updatedReply)
        updatedReplies.append(currentUpdatedReplies)


    if userType == "owner":
        return flask.redirect(flask.url_for("portal"))

    elif userType == "manager":
        SQL = """SELECT * FROM Managers
                    WHERE manager_id = ?"""
        userName = (EXECUTE(SQL, values=(userID,), getResult=True))[0][1]

        SQL = """SELECT * FROM Employees"""
        employees = EXECUTE(SQL, getResult=True)

        employeeDictionary = {}
        for employee in employees:
            employeeDictionary[employee[0]] = employee[1]

        SQL = """SELECT * FROM DirectMessages
                    WHERE sender_id = ?"""
        directMessages = EXECUTE(SQL, values=(userID,), getResult=True)
        updatedDirectMessages = []
        for message in directMessages:
            SQL = """SELECT * FROM ReplyMessage
                        WHERE message_id = ?"""
            reply = EXECUTE(SQL, values=(message[0],), getResult=True)
            if reply:
                currentReply = (employeeDictionary.get(message[2]), reply[0][2])
            else:
                currentReply = ()
            currentMessage = (employeeDictionary.get(message[2]), message[3])
            currentMessageContainer = [currentMessage, currentReply]
            updatedDirectMessages.append(currentMessageContainer)


        if flask.request.method == "GET":
            return flask.render_template("chat.html", usertype=userType, businessname=businessName, groupmessages=updatedMessages, groupreplies=updatedReplies, username=userName, employees=employees, directMessages=updatedDirectMessages, leftbar=True, rightbar=True)

        action = flask.request.form.get("action")

        if action == "messageGroup":
            message = flask.request.form.get("message")
            if message:
                SQL = """INSERT INTO GroupMessages (sender_id, message)
                            VALUES (?, ?)"""
                EXECUTE(SQL, values=(accountID, message))
                return flask.redirect(flask.url_for("chat"))
            else:
                flask.flash("Please Enter Message!")
                return flask.render_template("chat.html", usertype=userType, businessname=businessName, groupmessages=updatedMessages, groupreplies=updatedReplies, username=userName, employees=employees, directMessages=updatedDirectMessages, leftbar=True, rightbar=True)

        elif action == "replyGroup":
            message = flask.request.form.get("reply")
            messageID = flask.request.form.get("messageID")
            if message:
                SQL = """INSERT INTO GroupMessages (sender_id, reply_id, message)
                            VALUES (?, ?, ?)"""
                EXECUTE(SQL, values=(accountID, messageID, message))
                return flask.redirect(flask.url_for("chat"))
            else:
                flask.flash("Please Enter Message!")
                return flask.render_template("chat.html", usertype=userType, businessname=businessName, groupmessages=updatedMessages, groupreplies=updatedReplies, username=userName, employees=employees, directMessages=updatedDirectMessages, leftbar=True, rightbar=True)

        elif action == "directSend":
            employee = flask.request.form.get("selectEmployee")
            message = flask.request.form.get("directMessage")
            if employee and message:
                SQL = """INSERT INTO DirectMessages (sender_id, receiver_id, message)
                            VALUES (?, ?, ?)"""
                EXECUTE(SQL, values=(userID, employee, message))
                return flask.redirect(flask.url_for("chat"))
            else:
                flask.flash("Please Enter Message!")
                return flask.render_template("chat.html", usertype=userType, businessname=businessName, groupmessages=updatedMessages, groupreplies=updatedReplies, username=userName, employees=employees, directMessages=updatedDirectMessages, leftbar=True, rightbar=True)

    elif userType == "employee":
        SQL = """SELECT * FROM Employees
                    WHERE employee_id = ?"""
        userName = (EXECUTE(SQL, values=(userID,), getResult=True))[0][1]

        SQL = """SELECT * FROM Managers"""
        managers = EXECUTE(SQL, getResult=True)

        managerDictionary = {}
        for manager in managers:
            managerDictionary[manager[0]] = manager[1]

        SQL = """SELECT * FROM DirectMessages
                    WHERE receiver_id = ?"""
        directMessages = EXECUTE(SQL, values=(userID,), getResult=True)
        updatedDirectMessages = []
        for message in directMessages:
            SQL = """SELECT * FROM ReplyMessage
                        WHERE message_id = ?"""
            reply = EXECUTE(SQL, values=(message[0],), getResult=True)
            print(reply)
            if reply:
                continue
            else:
                currentMessage = (message[0], managerDictionary.get(message[1]), message[3])
                updatedDirectMessages.append(currentMessage)
        print(updatedDirectMessages)


        if flask.request.method == "GET":
            return flask.render_template("chat.html", usertype=userType, businessname=businessName, groupmessages=updatedMessages, groupreplies=updatedReplies, directmessages=updatedDirectMessages, username=userName, rightbar=True)

        action = flask.request.form.get("action")

        if action == "messageGroup":
            message = flask.request.form.get("message")
            if message:
                SQL = """INSERT INTO GroupMessages (sender_id, message)
                            VALUES (?, ?)"""
                EXECUTE(SQL, values=(accountID, message))
                return flask.redirect(flask.url_for("chat"))
            else:
                flask.flash("Please Enter Message!")
                return flask.render_template("chat.html", usertype=userType, businessname=businessName, groupmessages=updatedMessages, groupreplies=updatedReplies, directmessages=updatedDirectMessages, username=userName, rightbar=True)

        elif action == "replyGroup":
            message = flask.request.form.get("reply")
            messageID = flask.request.form.get("messageID")
            if message:
                SQL = """INSERT INTO GroupMessages (sender_id, reply_id, message)
                            VALUES (?, ?, ?)"""
                EXECUTE(SQL, values=(accountID, messageID, message))
                return flask.redirect(flask.url_for("chat"))
            else:
                flask.flash("Please Enter Message!")
                return flask.render_template("chat.html", usertype=userType, businessname=businessName, groupmessages=updatedMessages, groupreplies=updatedReplies, directmessages=updatedDirectMessages, username=userName, rightbar=True)

        elif action == "directEmployeeSend":
            message = flask.request.form.get("directEmployeeMessage")
            messageID = flask.request.form.get("directMessageID")
            print(messageID, message)
            if message:
                SQL = """INSERT INTO ReplyMessage (message_id, message)
                            VALUES (?, ?)"""
                EXECUTE(SQL, values=(messageID, message))
                return flask.redirect(flask.url_for("chat"))
            else:
                flask.flash("Please Enter Message!")
                return flask.render_template("chat.html", usertype=userType, businessname=businessName, groupmessages=updatedMessages, groupreplies=updatedReplies, directmessages=updatedDirectMessages, username=userName, rightbar=True)

        else:
            return flask.render_template("chat.html", usertype=userType, businessname=businessName, groupmessages=updatedMessages, groupreplies=updatedReplies, directmessages=updatedDirectMessages, username=userName, rightbar=True)

# Holidaybooking route, implements the holiday booking feature designed from the research #
@app.route("/portal/holidaybooking", methods=["GET", "POST"])
def holidaybooking():
    if checkLogged() == False:
        return flask.redirect(flask.url_for("index"))

    SQL = """SELECT * FROM Business"""
    businessName = EXECUTE(SQL, getResult=True)[0][0]
    userID = flask.session.get("UserID")
    userType = flask.session.get("UserType")

    if userType == "owner":
        return flask.redirect(flask.url_for("portal"))

    elif userType == "manager":
        SQL = """SELECT * FROM Employees"""
        foundEmployees = EXECUTE(SQL, getResult=True)
        employeeDictionary = {}
        for employee in foundEmployees:
            employeeDictionary[employee[0]] = employee[1]

        SQL = """SELECT * FROM HolidayRequest
                    WHERE accepted = 0"""
        foundHolidays = EXECUTE(SQL, getResult=True)
        updatedHolidays = []
        for holiday in foundHolidays:
            updatedHolidays.append((holiday[0], employeeDictionary.get(holiday[1]), holiday[2], holiday[3], holiday[4]))

        if flask.request.method == "GET":
            return flask.render_template("holidaybooking.html", usertype=userType, businessname=businessName, employees=foundEmployees, holidayrequests=updatedHolidays, rightbar=True)

        action = flask.request.form.get("action")
        if action == "approveRequest":
            requestID = flask.request.form.get("requestID")
            SQL = """UPDATE HolidayRequest
                        SET accepted = 1
                        WHERE holiday_request_id = ?"""
            EXECUTE(SQL, values=(requestID,))

            SQL = """SELECT * FROM HolidayRequest
                        WHERE holiday_request_id = ?"""
            foundRequest = EXECUTE(SQL, values=(requestID,), getResult=True)
            print(foundRequest)

            startdate = (datetime.datetime.strptime(foundRequest[0][2], "%Y-%m-%d")).date()
            enddate = (datetime.datetime.strptime(foundRequest[0][3], "%Y-%m-%d")).date()

            while startdate <= enddate:
                SQL = """INSERT INTO Calendar (event_name, date)
                            VALUES (?, ?)"""
                eventname = employeeDictionary.get(foundRequest[0][1]) + " Time Off"
                print(eventname, startdate)
                EXECUTE(SQL, values=(eventname, startdate))
                startdate = startdate + datetime.timedelta(days=1)

            return flask.redirect(flask.url_for("holidaybooking"))
        elif action == "denyRequest":
            requestID = flask.request.form.get("requestID")
            reason = flask.request.form.get("denyReason")
            SQL = """DELETE FROM HolidayRequest
                        WHERE holiday_request_id = ?"""
            EXECUTE(SQL, values=(requestID,))
            return flask.redirect(flask.url_for("holidaybooking"))

        elif action == "addHoliday":
            employee = flask.request.form.get("employee")
            startDate = flask.request.form.get("addStartdate")
            endDate = flask.request.form.get("addEnddate")
            if startDate and endDate and employee:

                startDate = datetime.datetime.strptime(startDate, "%Y-%m-%d")
                endDate = datetime.datetime.strptime(endDate, "%Y-%m-%d")

                if endDate < startDate:
                    flask.flash("End date must be equal or after start date!")
                    return flask.render_template("holidaybooking.html", usertype=userType, businessname=businessName, employees=foundEmployees, holidayrequests=updatedHolidays, rightbar=True)

                startDate = datetime.datetime.strftime(startDate, "%Y-%m-%d")
                endDate = datetime.datetime.strftime(endDate, "%Y-%m-%d")

                SQL = """INSERT INTO HolidayRequest (employee_id, startdate, enddate, accepted)
                            VALUES (?, ?, ?, ?)"""
                EXECUTE(SQL, values=(employee, startDate, endDate, 1))
                return flask.redirect(flask.url_for("holidaybooking"))
            else:
                flask.flash("Please Fill All Fields!")
                return flask.render_template("holidaybooking.html", usertype=userType, businessname=businessName, employees=foundEmployees, holidayrequests=updatedHolidays, rightbar=True)
        
        else:
            return flask.render_template("holidaybooking.html", usertype=userType, businessname=businessName, employees=foundEmployees, holidayrequests=updatedHolidays, rightbar=True)

    elif userType == "employee":
        if flask.request.method == "GET":
            return flask.render_template("holidaybooking.html", usertype=userType, businessname=businessName)

        action = flask.request.form.get("action")
        if action == "submitRequest":
            startDate = flask.request.form.get("startdate")
            endDate = flask.request.form.get("enddate")
            reason = flask.request.form.get("reason")
            print(startDate, endDate, reason)
            if startDate and endDate and reason:

                startDate = datetime.datetime.strptime(startDate, "%Y-%m-%d")
                endDate = datetime.datetime.strptime(endDate, "%Y-%m-%d")

                if endDate < startDate:
                    flask.flash("End date must be equal or after start date!")
                    return flask.render_template("holidaybooking.html", usertype=userType, businessname=businessName)

                startDate = datetime.datetime.strftime(startDate, "%Y-%m-%d")
                endDate = datetime.datetime.strftime(endDate, "%Y-%m-%d")

                SQL = """INSERT INTO HolidayRequest (employee_id, startdate, enddate, reason, accepted)
                            VALUES (?, ?, ?, ?, ?)"""
                EXECUTE(SQL, values=(userID, startDate, endDate, reason, 0))
                return flask.redirect(flask.url_for("holidaybooking"))
            else:
                flask.flash("Please Fill All Fields!")
                return flask.render_template("holidaybooking.html", usertype=userType, businessname=businessName)
        else:
            return flask.render_template("holidaybooking.html", usertype=userType, businessname=businessName)

# Sales route, implements the sales feature designed from the research #
@app.route("/portal/sales", methods=["GET", "POST"])
def sales():
    if checkLogged() == False:
        return flask.redirect(flask.url_for("index"))

    SQL = """SELECT * FROM Business"""
    businessName = EXECUTE(SQL, getResult=True)[0][0]
    userID = flask.session.get("UserID")
    userType = flask.session.get("UserType")

    if userType == "owner":
        today = datetime.datetime.now()
        position = today.weekday()
        startdate = (today - datetime.timedelta(days=position)).date()
        enddate = (today + datetime.timedelta(days=6-position)).date()
        SQL = """SELECT * FROM Products"""
        foundProducts = EXECUTE(SQL, getResult=True)
        productTotals = []

        for product in foundProducts:
            currentdate = startdate
            currentTotal = 0
            while currentdate <= enddate:
                SQL = """SELECT * FROM Sales
                            WHERE date = ?
                            AND product_id = ?"""
                foundSales = EXECUTE(SQL, values=(currentdate.strftime("%Y-%m-%d"), product[0]), getResult=True)
                for sale in foundSales:
                    currentTotal += sale[4]
                currentdate = currentdate + datetime.timedelta(days=1)
            productTotals.append((product[0], product[1], currentTotal))
        SQL = """SELECT * FROM Employees"""
        foundEmployees = EXECUTE(SQL, getResult=True)
        employeeTotals = []

        for employee in foundEmployees:
            currentdate = startdate
            currentTotal = 0
            while currentdate <= enddate:
                SQL = """SELECT * FROM Sales
                            WHERE date = ?
                            AND employee_id = ?"""
                foundSales = EXECUTE(SQL, values=(currentdate.strftime("%Y-%m-%d"), employee[0]), getResult=True)
                for sale in foundSales:
                    currentTotal += 1
                currentdate = currentdate + datetime.timedelta(days=1)
            employeeTotals.append((employee[1], currentTotal))

        totalMade = 0
        if (len(productTotals) != 0):
            for product in productTotals:
                SQL = """SELECT * FROM ProductCosts
                            WHERE product_id = ?"""
                cost = EXECUTE(SQL, values=(product[0],), getResult=True)
                totalMade = round(totalMade + (cost[0][1]*product[2]), 2)

        if flask.request.method == "GET":
            return flask.render_template("sales.html", usertype=userType, businessname=businessName, date=startdate, products=productTotals, employeesSales=employeeTotals, total=totalMade)

        action = flask.request.form.get("action")

        if action == "previous":
            date = flask.request.form.get("date")
            foundDate = datetime.datetime.strptime(date, "%Y-%m-%d")
            newDate = (foundDate - datetime.timedelta(days=7)).date()
            newEndDate = (foundDate - datetime.timedelta(days=1)).date()
            SQL = """SELECT * FROM Products"""
            foundProducts = EXECUTE(SQL, getResult=True)
            productTotals = []

            for product in foundProducts:
                currentdate = newDate
                currentTotal = 0
                while currentdate <= newEndDate:
                    SQL = """SELECT * FROM Sales
                                WHERE date = ?
                                AND product_id = ?"""
                    foundSales = EXECUTE(SQL, values=(currentdate.strftime("%Y-%m-%d"), product[0]), getResult=True)
                    for sale in foundSales:
                        currentTotal += sale[4]
                    currentdate = currentdate + datetime.timedelta(days=1)
                productTotals.append((product[0], product[1], currentTotal))
            SQL = """SELECT * FROM Employees"""
            foundEmployees = EXECUTE(SQL, getResult=True)
            employeeTotals = []

            for employee in foundEmployees:
                currentdate = newDate
                currentTotal = 0
                while currentdate <= newEndDate:
                    SQL = """SELECT * FROM Sales
                                WHERE date = ?
                                AND employee_id = ?"""
                    foundSales = EXECUTE(SQL, values=(currentdate.strftime("%Y-%m-%d"), employee[0]), getResult=True)
                    for sale in foundSales:
                        currentTotal += 1
                    currentdate = currentdate + datetime.timedelta(days=1)
                employeeTotals.append((employee[1], currentTotal))

            totalMade = 0
            for product in productTotals:
                SQL = """SELECT * FROM ProductCosts
                            WHERE product_id = ?"""
                cost = EXECUTE(SQL, values=(product[0],), getResult=True)
                totalMade = round(totalMade + (cost[0][1]*product[2]), 2)

            return flask.render_template("sales.html", usertype=userType, businessname=businessName, date=newDate, products=productTotals, employeesSales=employeeTotals, total=totalMade)

        elif action == "next":
            date = flask.request.form.get("date")
            foundDate = datetime.datetime.strptime(date, "%Y-%m-%d")
            newDate = (foundDate + datetime.timedelta(days=7)).date()
            newEndDate = (foundDate + datetime.timedelta(days=13)).date()
            print(foundDate, newDate, newEndDate)
            SQL = """SELECT * FROM Products"""
            foundProducts = EXECUTE(SQL, getResult=True)
            productTotals = []

            for product in foundProducts:
                currentdate = newDate
                currentTotal = 0
                while currentdate <= newEndDate:
                    SQL = """SELECT * FROM Sales
                                WHERE date = ?
                                AND product_id = ?"""
                    foundSales = EXECUTE(SQL, values=(currentdate.strftime("%Y-%m-%d"), product[0]), getResult=True)
                    for sale in foundSales:
                        currentTotal += sale[4]
                    currentdate = currentdate + datetime.timedelta(days=1)
                productTotals.append((product[0], product[1], currentTotal))
            SQL = """SELECT * FROM Employees"""
            foundEmployees = EXECUTE(SQL, getResult=True)
            employeeTotals = []

            for employee in foundEmployees:
                currentdate = newDate
                currentTotal = 0
                while currentdate <= newEndDate:
                    SQL = """SELECT * FROM Sales
                                WHERE date = ?
                                AND employee_id = ?"""
                    foundSales = EXECUTE(SQL, values=(currentdate.strftime("%Y-%m-%d"), employee[0]), getResult=True)
                    for sale in foundSales:
                        currentTotal += 1
                    currentdate = currentdate + datetime.timedelta(days=1)
                employeeTotals.append((employee[1], currentTotal))

            totalMade = 0
            for product in productTotals:
                SQL = """SELECT * FROM ProductCosts
                            WHERE product_id = ?"""
                cost = EXECUTE(SQL, values=(product[0],), getResult=True)
                totalMade = round(totalMade + (cost[0][1]*product[2]), 2)

            return flask.render_template("sales.html", usertype=userType, businessname=businessName, date=newDate, products=productTotals, employeesSales=employeeTotals, total=totalMade)
    else:
        return flask.redirect(flask.url_for("portal"))

# Sheets route, implements the timesheets/paysheets feature designed from the research #
@app.route("/portal/sheets", methods=["GET", "POST"])
def sheets():
    if checkLogged() == False:
        return flask.redirect(flask.url_for("index"))

    SQL = """SELECT * FROM Business"""
    businessName = EXECUTE(SQL, getResult=True)[0][0]
    userID = flask.session.get("UserID")
    userType = flask.session.get("UserType")

    if userType == "owner":
        return flask.redirect(flask.url_for("portal"))

    elif userType == "manager":
        SQL = """SELECT * FROM Employees"""
        foundEmployees = EXECUTE(SQL, getResult=True)
        employeeDictionary = {}
        for employee in foundEmployees:
            employeeDictionary[employee[0]] = employee[1]

        if flask.request.method == "GET":
            return flask.render_template("sheets.html", usertype=userType, businessname=businessName, employees=foundEmployees)

        action = flask.request.form.get("action")
        if action == "download":
            employee = flask.request.form.get("employee")
            month = int(flask.request.form.get("month"))
            year = datetime.datetime.now().year
            startdate = datetime.date(year, month, 1)
            current = startdate
            timesheet = docx.Document()
            timesheet.add_heading(f"Time Sheet for {employeeDictionary.get(int(employee))}: Month Starting {startdate.strftime("%Y-%m-%d")}", 0)
            finalTotalHours = 0
            while startdate.month == current.month:
                timesheet.add_paragraph(f"\n\n{current.strftime("%Y-%m-%d")}")
                SQL = """SELECT * FROM Shifts
                            WHERE employee_id = ?
                            AND date = ?"""
                shifts = EXECUTE(SQL, values=(employee, current.strftime("%Y-%m-%d")), getResult=True)#
                shiftText = "\n"
                hoursTotal = 0
                for shift in shifts:
                    shiftText = shiftText + shift[3] + " - " + shift[4] + "\n"
                    starttime = datetime.datetime.strptime(shift[3],"%H:%M")
                    endtime = datetime.datetime.strptime(shift[4], "%H:%M")
                    hours = endtime - starttime
                    hours = hours.seconds / 3600
                    hoursTotal = hoursTotal + hours
                shiftText = shiftText + "Total Hours: " + str(hoursTotal)
                timesheet.add_paragraph(shiftText)

                finalTotalHours = finalTotalHours + hoursTotal
                current = current + datetime.timedelta(days=1)
            timesheet.add_paragraph(f"\n\n\nMonth Total Hours Worked: {finalTotalHours}")
            filename = f"{startdate}{employee}.docx"
            filepath = os.path.join(app.config["TIME_SHEETS"], filename)
            timesheet.save(filepath)
            time.sleep(1)
            return flask.send_from_directory(app.config["TIME_SHEETS"], filename, as_attachment=True)
            # return flask.redirect(flask.url_for("sheets"))
        else:
            return flask.render_template("sheets.html", usertype=userType, businessname=businessName, employees=foundEmployees)



    elif userType == "employee":

        if flask.request.method == "GET":
            return flask.render_template("sheets.html", usertype=userType, businessname=businessName)

        action = flask.request.form.get("action")
        if action == "download":
            month = int(flask.request.form.get("month"))
            year = datetime.datetime.now().year
            startdate = datetime.date(year, month, 1)
            current = startdate
            timesheet = docx.Document()
            timesheet.add_heading(f"Paysheet for Month Starting {startdate.strftime("%Y-%m-%d")}", 0)
            hoursWorked = 0
            
            SQL = """SELECT * FROM PayRates
                        WHERE employee_id = ?"""
            payrate = EXECUTE(SQL, values=(userID,), getResult=True)
            while startdate.month == current.month:
                SQL = """SELECT * FROM Shifts
                            WHERE employee_id = ?
                            AND date = ?"""
                shifts = EXECUTE(SQL, values=(userID, current.strftime("%Y-%m-%d")), getResult=True)#
                hoursTotal = 0
                for shift in shifts:
                    starttime = datetime.datetime.strptime(shift[3],"%H:%M")
                    endtime = datetime.datetime.strptime(shift[4], "%H:%M")
                    hours = endtime - starttime
                    hours = hours.seconds / 3600
                    hoursTotal = hoursTotal + hours

                hoursWorked = hoursWorked + hoursTotal
                current = current + datetime.timedelta(days=1)
            timesheet.add_paragraph(f"\n\n\nHours Worked: {hoursWorked}")
            payrate = payrate[0][2]
            payAmount = round((hoursWorked*payrate), 2)
            print(payAmount)
            timesheet.add_paragraph(f"\n\n\nTotal Pay: {payAmount}")
            filename = f"{startdate}{userID}.docx"
            filepath = os.path.join(app.config["PAY_SHEETS"], filename)
            timesheet.save(filepath)
            time.sleep(1)
            return flask.send_from_directory(app.config["PAY_SHEETS"], filename, as_attachment=True)
            # return flask.redirect(flask.url_for("sheets"))

# Stock route, implements the stock management feature designed from the research #
@app.route("/portal/stock", methods=["GET", "POST"])
def stock():
    if checkLogged() == False:
        return flask.redirect(flask.url_for("index"))

    SQL = """SELECT * FROM Business"""
    businessName = EXECUTE(SQL, getResult=True)[0][0]
    userID = flask.session.get("UserID")
    userType = flask.session.get("UserType")

    if userType == "owner" or userType == "employee":
        return flask.redirect(flask.url_for("portal"))
    elif userType == "manager":

        SQL = """SELECT * FROM Products"""
        foundStock = EXECUTE(SQL, getResult=True)

        alerts = []
        forecast = [[],[]]
        for stock in foundStock:
            if stock[4]:
                if stock[2] < stock[4]:
                    alerts.append((stock[1], stock[4]))

            SQL = """SELECT * FROM Sales
                        WHERE product_id = ?"""
            foundSales = EXECUTE(SQL, values=(stock[0],), getResult=True)
            totalSold = 0
            for sale in foundSales:
                totalSold += sale[4]
            averageDailySold = round(totalSold/7)
            averageWeeklySold = averageDailySold*7
            week1Forecast = stock[2] - averageWeeklySold
            week2Forecast = stock[2] - (2 * averageWeeklySold)
            if week1Forecast >= 0:
                forecast[0].append((stock[1],week1Forecast))
            else:
                forecast[0].append((stock[1],0))

            if week2Forecast >= 0:
                forecast[1].append((stock[1],week2Forecast))
            else:
                forecast[1].append((stock[1],0))

        

        if flask.request.method == "GET":
            return flask.render_template("stock.html", usertype=userType, businessname=businessName, stocks=foundStock, alerts=alerts, forecast=forecast, leftbar=True, rightbar=True)
        action = flask.request.form.get("action")

        if action == "createStock":
            name = flask.request.form.get("name")
            amount = flask.request.form.get("amount")
            date = flask.request.form.get("date")
            minAmount = flask.request.form.get("minAmount")
            if name and amount and minAmount:
                amount = int(amount)
                minAmount = int(minAmount)
                if amount < 1 or minAmount < 1:
                    flask.flash("Amounts Should Be Greater than 0!")
                    return flask.render_template("stock.html", usertype=userType, businessname=businessName, stocks=foundStock, alerts=alerts, forecast=forecast, leftbar=True, rightbar=True)
                else:
                    if date:
                        SQL = """INSERT INTO Products (product_name, amount_stocked, perish_date, min_amount)
                                    VALUES (?, ?, ?, ?)"""
                        EXECUTE(SQL, values=(name, amount, date, minAmount))
                        return flask.redirect(flask.url_for("stock"))
                    else:
                        SQL = """INSERT INTO Products (product_name, amount_stocked, min_amount)
                                    VALUES (?, ?, ?)"""
                        EXECUTE(SQL, values=(name, amount, minAmount))
                        return flask.redirect(flask.url_for("stock"))
            else:
                flask.flash("Please Fill All Fields!")
                return flask.render_template("stock.html", usertype=userType, businessname=businessName, stocks=foundStock, alerts=alerts, forecast=forecast, leftbar=True, rightbar=True)
        elif action == "increase":
            stockID = flask.request.form.get("stockID")
            currentAmount = int(flask.request.form.get("currentAmount"))
            value = flask.request.form.get("value")

            if value:
                value = int(value)
                if value < 1:
                    flask.flash("Increase Amount Should Be Greater Than 0!")
                    return flask.render_template("stock.html", usertype=userType, businessname=businessName, stocks=foundStock, alerts=alerts, forecast=forecast, leftbar=True, rightbar=True)
                else:
                    newAmount = currentAmount + value

                    SQL = """UPDATE Products
                                SET amount_stocked = ?
                                WHERE product_id = ?"""
                    EXECUTE(SQL, values=(newAmount, stockID))
                    return flask.redirect(flask.url_for("stock"))
            else:
                flask.flash("Increase Amount Should Be Greater Than 0!")
                return flask.render_template("stock.html", usertype=userType, businessname=businessName, stocks=foundStock, alerts=alerts, forecast=forecast, leftbar=True, rightbar=True)

        elif action == "decrease":
            stockID = flask.request.form.get("stockID")
            currentAmount = int(flask.request.form.get("currentAmount"))
            value = flask.request.form.get("value")

            if value:
                value = int(value)
                if value < 1:
                    flask.flash("Decrease Amount Should Be Greater Than 0!")
                    return flask.render_template("stock.html", usertype=userType, businessname=businessName, stocks=foundStock, alerts=alerts, forecast=forecast, leftbar=True, rightbar=True)
                else:
                    newAmount = currentAmount - value
                    if newAmount < 0:
                        newAmount = 0

                    SQL = """UPDATE Products
                                SET amount_stocked = ?
                                WHERE product_id = ?"""
                    EXECUTE(SQL, values=(newAmount, stockID))
                    return flask.redirect(flask.url_for("stock"))
            else:
                flask.flash("Decrease Amount Should Be Greater Than 0!")
                return flask.render_template("stock.html", usertype=userType, businessname=businessName, stocks=foundStock, alerts=alerts, forecast=forecast, leftbar=True, rightbar=True)
        
        elif action == "delete":
            stockID = flask.request.form.get("stockID")
            SQL = """DELETE FROM Products
                        WHERE product_id = ?"""
            EXECUTE(SQL, values=(stockID,))
            return flask.redirect(flask.url_for("stock"))

        else:
            return flask.render_template("stock.html", usertype=userType, businessname=businessName, stocks=foundStock, alerts=alerts, forecast=forecast, leftbar=True, rightbar=True)

# Calendar route, implements the calendar feature designed from the research #
@app.route("/portal/calendar", methods=["GET", "POST"])
def calendar():
    if checkLogged() == False:
        return flask.redirect(flask.url_for("index"))

    SQL = """SELECT * FROM Business"""
    businessName = EXECUTE(SQL, getResult=True)[0][0]
    userID = flask.session.get("UserID")
    userType = flask.session.get("UserType")

    if userType == "owner":
        return flask.redirect(flask.url_for("portal"))
    
    elif userType == "manager":
        today = datetime.datetime.now()
        startdate = today.date()
        foundEvents = []
        date = []
        for i in range(3):
            date.append(startdate)
            SQL = """SELECT * FROM Calendar
                        WHERE date = ?"""
            currentEvents = EXECUTE(SQL, values=(startdate,), getResult=True)
            foundEvents.append(currentEvents)
            startdate = startdate + datetime.timedelta(days=1)
        if flask.request.method == "GET":
            return flask.render_template("calendar.html", usertype=userType, businessname=businessName, date=date, events=foundEvents, rightbar=True)
        action = flask.request.form.get("action")

        if action == "addEvent":
            eventName = flask.request.form.get("eventName")
            startdate = flask.request.form.get("startdate")
            enddate = flask.request.form.get("enddate")

            if eventName and startdate and enddate:

                newStartDate = (datetime.datetime.strptime(startdate, "%Y-%m-%d")).date()
                newEndDate = (datetime.datetime.strptime(enddate, "%Y-%m-%d")).date()
                while newStartDate <= newEndDate:
                    SQL = """INSERT INTO Calendar (event_name, date)
                                VALUES (?, ?)"""
                    EXECUTE(SQL, values=(eventName, datetime.datetime.strftime(newStartDate, "%Y-%m-%d")))
                    newStartDate = newStartDate + datetime.timedelta(days=1)
                return flask.redirect(flask.url_for("calendar"))
            else:
                flask.flash("Please Fill All Event Fields!")
                return flask.render_template("calendar.html", usertype=userType, businessname=businessName, date=date, events=foundEvents, rightbar=True)

        elif action == "deleteEvent":
            eventID = flask.request.form.get("eventID")
            SQL = """DELETE FROM Calendar
                        WHERE calendar_id = ?"""
            EXECUTE(SQL, values=(eventID,))
            return flask.redirect(flask.url_for("calendar"))

        elif action == "next":
            currentDate = (datetime.datetime.strptime(flask.request.form.get("currentdate"), "%Y-%m-%d")).date()
            startdate = currentDate + datetime.timedelta(days=3)
            foundEvents = []
            date = []
            for i in range(3):
                date.append(startdate)
                SQL = """SELECT * FROM Calendar
                            WHERE date = ?"""
                currentEvents = EXECUTE(SQL, values=(startdate,), getResult=True)
                foundEvents.append(currentEvents)
                startdate = startdate + datetime.timedelta(days=1)
            return flask.render_template("calendar.html", usertype=userType, businessname=businessName, date=date, events=foundEvents, rightbar=True)

        elif action == "previous":
            currentDate = (datetime.datetime.strptime(flask.request.form.get("currentdate"), "%Y-%m-%d")).date()
            startdate = currentDate - datetime.timedelta(days=3)
            foundEvents = []
            date = []
            for i in range(3):
                date.append(startdate)
                SQL = """SELECT * FROM Calendar
                            WHERE date = ?"""
                currentEvents = EXECUTE(SQL, values=(startdate,), getResult=True)
                foundEvents.append(currentEvents)
                startdate = startdate + datetime.timedelta(days=1)
            return flask.render_template("calendar.html", usertype=userType, businessname=businessName, date=date, events=foundEvents, rightbar=True)
        else:
            return flask.render_template("calendar.html", usertype=userType, businessname=businessName, date=date, events=foundEvents, rightbar=True)

    elif userType == "employee":
        today = datetime.datetime.now()
        startdate = today.date()
        foundEvents = []
        date = []
        for i in range(3):
            date.append(startdate)
            SQL = """SELECT * FROM Calendar
                        WHERE date = ?"""
            currentEvents = EXECUTE(SQL, values=(startdate,), getResult=True)
            foundEvents.append(currentEvents)
            startdate = startdate + datetime.timedelta(days=1)
        if flask.request.method == "GET":
            return flask.render_template("calendar.html", usertype=userType, businessname=businessName, date=date, events=foundEvents)
        action = flask.request.form.get("action")

        if action == "next":
            currentDate = (datetime.datetime.strptime(flask.request.form.get("currentdate"), "%Y-%m-%d")).date()
            startdate = currentDate + datetime.timedelta(days=3)
            foundEvents = []
            date = []
            for i in range(3):
                date.append(startdate)
                SQL = """SELECT * FROM Calendar
                            WHERE date = ?"""
                currentEvents = EXECUTE(SQL, values=(startdate,), getResult=True)
                foundEvents.append(currentEvents)
                startdate = startdate + datetime.timedelta(days=1)
            return flask.render_template("calendar.html", usertype=userType, businessname=businessName, date=date, events=foundEvents)

        elif action == "previous":
            currentDate = (datetime.datetime.strptime(flask.request.form.get("currentdate"), "%Y-%m-%d")).date()
            startdate = currentDate - datetime.timedelta(days=3)
            foundEvents = []
            date = []
            for i in range(3):
                date.append(startdate)
                SQL = """SELECT * FROM Calendar
                            WHERE date = ?"""
                currentEvents = EXECUTE(SQL, values=(startdate,), getResult=True)
                foundEvents.append(currentEvents)
                startdate = startdate + datetime.timedelta(days=1)
            return flask.render_template("calendar.html", usertype=userType, businessname=businessName, date=date, events=foundEvents)
        else:
            return flask.render_template("calendar.html", usertype=userType, businessname=businessName, date=date, events=foundEvents)



# Route for rendering a requested image #
@app.route("/images/<name>")
def getImage(name):
    return flask.send_from_directory(app.config["UPLOAD_FOLDER"], name)

# Route for clearing session (logging out)
@app.route("/logout")
def logout():
    flask.session.clear()
    return flask.redirect(flask.url_for("index"))

if __name__ == "__main__":
    app.run(debug=True)